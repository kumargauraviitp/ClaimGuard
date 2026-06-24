import os
import uuid
import markdown
from datetime import datetime
from docx import Document
from sqlalchemy.orm import Session
from app.models import InvestigationReport, ExportQueue

EXPORT_DIR = "storage/exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

class ExportService:
    def __init__(self, db: Session):
        self.db = db

    def _render_markdown(self, report: InvestigationReport) -> str:
        # Mock rendering report sections into markdown
        md = f"# Investigation Report\n\n**Claim ID**: {report.claim_id}\n\n"
        if report.executive_summary:
            md += f"## Executive Summary\n{report.executive_summary}\n\n"
        if report.confidence_analysis:
            md += f"## AI Confidence Explanation\nOverall Confidence: {report.confidence_analysis.get('overall_confidence')}\n\n"
        return md

    def _generate_pdf(self, markdown_text: str, file_path: str, is_draft: bool):
        try:
            from weasyprint import HTML, CSS
        except ImportError as e:
            raise Exception("PDF generation is not available. Missing system dependencies like Pango or Gtk. " + str(e))
            
        html_string = markdown.markdown(markdown_text)
        
        # Simple Watermarking CSS
        watermark_css = """
        @page {
            @background {
                content: "DRAFT";
                color: rgba(200, 200, 200, 0.3);
                font-size: 150pt;
                transform: rotate(-45deg);
                position: absolute;
                top: 30%;
                left: 10%;
            }
        }
        """ if is_draft else ""
        
        HTML(string=html_string).write_pdf(file_path, stylesheets=[CSS(string=watermark_css)] if is_draft else [])

    def _generate_docx(self, markdown_text: str, file_path: str):
        # Extremely simplified Markdown to DOCX mapping
        doc = Document()
        lines = markdown_text.split("\n")
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            else:
                if line.strip():
                    doc.add_paragraph(line)
        doc.save(file_path)

    def process_export_job(self, job_id: uuid.UUID):
        job = self.db.query(ExportQueue).filter(ExportQueue.id == job_id).first()
        if not job:
            return
        
        try:
            job.status = "processing"
            self.db.commit()

            report = self.db.query(InvestigationReport).filter(InvestigationReport.id == job.report_id).first()
            if not report:
                raise Exception("Report not found")

            md_text = self._render_markdown(report)
            is_draft = not report.locked_after_approval
            
            file_name = f"{report.claim_id}_{job.format}_{uuid.uuid4()}.{job.format}"
            file_path = os.path.join(EXPORT_DIR, file_name)

            if job.format == "pdf":
                self._generate_pdf(md_text, file_path, is_draft)
            elif job.format == "docx":
                self._generate_docx(md_text, file_path)
            elif job.format == "html":
                html_string = markdown.markdown(md_text)
                with open(file_path, "w") as f:
                    f.write(html_string)
            else:
                raise Exception("Unsupported format")

            job.file_path = file_path
            job.status = "completed"
            job.completed_at = datetime.utcnow()

        except Exception as e:
            job.status = "failed"
            job.error_message = str(e)
        finally:
            self.db.commit()
